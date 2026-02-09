from dotenv import load_dotenv
import os
from langgraph.graph import StateGraph, START, END
from typing import TypedDict, Literal, Annotated
from pydantic import BaseModel, Field
from langchain_openai import ChatOpenAI
import operator
import streamlit as st
from docx import Document
from pages.model_selection import selecting_model
import io
import time


if 'register' not in st.session_state:
    st.session_state['register'] = False

elif st.session_state['register'] == True:
    load_dotenv(
        dotenv_path=r"my_api_key.env")
    os.environ['OPENAI_API_KEY'] = os.getenv("openai_api")

    # ------------------model selection-------------------
    evaluate_model = ChatOpenAI(model="gpt-3.5-turbo")
    optimized_model = ChatOpenAI(model="gpt-5-nano-2025-08-07")

    st.title("📑Outline Generation")
    st.space("small")
    models = selecting_model()
    model = ChatOpenAI(model=models)

    # ---------------------EVALUATING CLASS--------------------

    class evaluating(BaseModel):
        evaluate: Literal["approved", "not_approved"] = Field(
            description="evaluate the content")
        feedback:  str = Field(description="feedback for the content")

    # ------------------CREATED STATE -------------------------------

    class outlinestate(TypedDict):

        topic:   str = Field(description="topic for the material")
        outline:   str = Field(description="outline for the topic")
        evaluate:   Literal["approved", "not_approved"]
        feedback:   str = Field(description="feedback for the material")
        iteration:   int
        max_iteration:   int
        outline_history:   Annotated[list[str], operator.add]
        feedback_history:   Annotated[list[str], operator.add]

    struct_model = evaluate_model.with_structured_output(evaluating)

    def generate_outline(state: outlinestate):
        prompt = f"""
        create a descriptive outline on {state['topic']} explain each topic and sub topic from it,make the topic look bold , you have to replace the normal words with seo keywords where ever required
        include the topics nd sub-topics in points 
        you have to explore all the tutorial websites and generate the content.
        Must include all the topic and beginner friendly.
        """

        response = model.invoke(prompt).content

        return {'outline': response, 'outline_history': [response]}

    def evaluating_content(state: outlinestate):
        prompt = f"evaluate the content on the basis of the {state['outline']} and check if the content is to be approved or not."

        response = struct_model.invoke(prompt)

        return {"evaluate": response, 'feedback': response.feedback, 'feedback_history': response.feedback}

    def optimize(state: outlinestate):
        prompt = f"improve the content on the basis of the feedback  {state['feedback_history']}"

        response = optimized_model.invoke(prompt)
        iteration = state['iteration'] + 1

        return {'outline': response, 'iteration': iteration, 'outline_history': [response]}

    def evaluate_route(state: outlinestate):
        if state['evaluate'] == 'approved' and state['iteration'] >= state['max_iteration']:
            return "approved"
        else:
            return "not_approved"

    # --------------------------------------------------------------------------
    st.space("small")
    st.subheader("Upload Your File Below👇")
    st.caption("✅Make sure The File Must Be Docx File")
    read_file = st.file_uploader("choose a file")
    full_text = []

    # reading the plain text from docx
    if read_file is not None:
        doc = Document(read_file)

        for para in doc.paragraphs:
            full_text.append(para.text)
    else:
        st.warning("Please upload the file first")

    content = "\n".join(full_text)
    st.markdown(content)

    # ---------------------------------------------------------------------------------
    # creating a graph and compiling it
    graph = StateGraph(outlinestate)

    # adding the node for the graph
    graph.add_node("generate_outline", generate_outline)
    graph.add_node("evaluate", evaluating_content)
    graph.add_node("optimize", optimize)

    # adding the edges
    graph.add_edge(START, "generate_outline")
    graph.add_conditional_edges("evaluate", evaluate_route, {
                                "approved": END, "not_approved": 'optimize'})
    graph.add_edge("optimize", 'evaluate')

     # compiling graph
    
    initial_state = {
        "topic": content,
        "iteration": 1,
        "max_iteration ": 3
    }

    # taking user input and invoke
    generate = st.button(label="GENERATE OUTLINE", use_container_width=True)

    if generate:
        with st.spinner("⛷️Generating Response...."):
            time.sleep(5)
            workflow = graph.compile()
            final_output = workflow.invoke(initial_state)
            outline_output = final_output['outline']
            st.markdown(outline_output)

        # creating a session state to store the data
        st.session_state["outline_output"] = outline_output

    try:
        insert_outline = st.button("INSERT DATA INTO FILE", use_container_width=True)

        if insert_outline:
            outline_file_name = read_file.name.replace(".docx", "")
            if "outline_output" not in st.session_state:
                st.error("Please generate the outline first!!")

            else:
                # reading exixting file
                try:
                    doc = Document(f"{outline_file_name}_outline.docx")

                except Exception as e:
                    doc = Document()

                doc.add_paragraph(st.session_state['outline_output'])

                # save the doc
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                st.success("📝Data inserted successfully!!")

                # creating a download button to download the file
                st.download_button(
                    label="DOWNLOAD FILE",
                    data=bio.getvalue(),
                    file_name=f"{outline_file_name}_outline.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    icon=":material/download:",
                    width = 'stretch'
                )

    except Exception as e:
        st.error(f"{e}")

    final_material_button = st.button(
        "GENERATE MATERIAL", use_container_width=True)
    roadmap_button = st.button("GENERATE ROADMAP", use_container_width=True)

    if final_material_button:
        st.switch_page("pages/final_material.py")
    elif roadmap_button:
        st.switch_page("pages/road_map.py")


else:
    st.switch_page("password.py")


from langgraph.graph import StateGraph, START, END
from dotenv import load_dotenv
import os
from langchain_openai import ChatOpenAI
from pydantic import BaseModel, Field
import operator
from typing import TypedDict, Literal, Annotated
import streamlit as st
from docx import Document
from pages.model_selection import selecting_model
import io
import time 

if 'register' not in st.session_state:
    st.session_state['register'] = False

elif st.session_state['register'] == True:
    load_dotenv(dotenv_path=r"my_api_key.env")
    os.environ["OPENAI_API_KEY"] = os.getenv("openai_api")
    
  

    # using openai model
    # -----------------------model selection --------------------
    evaluate_model = ChatOpenAI(model="gpt-3.5-turbo")
    optimized_model = ChatOpenAI(model="gpt-5-nano-2025-08-07")
    st.title(" 👨‍💻Study Material Generation")
    st.space("small")
    
    models = selecting_model()
    model = ChatOpenAI(model=models)

    # ---------------------------------------------------------------
    # creating a pydantic model to evalaute the content

    class evalauting(BaseModel):
        # literals are fixed value that represents the constant data (unchangeable)
        evaluate:   Literal["approved", "not_approved"] = Field(
            description="evaluating the topic.")
        feedback:   str = Field(
            description="feedback for the content generated.")

    # giving the structured output to the model
    struct_model = evaluate_model.with_structured_output(evalauting)

    # creating a state

    class studystate(TypedDict):
        topic: str = Field(description="topic for the material")
        prompt : str = Field(description="prompt for generation")
        material: str = Field(description="study notes for the topic")
        evaluate: Literal["approved", "not_approved"]
        feedback: str = Field(description="feedback for the evaluation")
        iteration: int = Field(
            description="iterate the model to improve the content")
        max_iteration: int = Field(
            description="includes the max nuber of iteration")
        material_history: str = Annotated[list[str], operator.add]
        # annotated allows to add meta data to type hint.
        feedback_history: str = Annotated[list[str], operator.add]

    # creating a function to provide a prompt

    def study_material(state:  studystate):
        prompt = f"""You are an experienced professional educator and tutor.
            Write a structured, theory-focused educational Study Material on {state['topic']}.
            Based on the given prompt{state['prompt']} you have to genrate the final output.         
        """
        
        response = model.invoke(prompt).content
        return {"material": response, "material_history": [response] , 'prompt': response}

    def evaluate_material(state: studystate):
        prompt = f"you have to evaluate the material and improve the content on the basis of the  {state['material']} , on topic {state['topic']}"

        response = struct_model.invoke(prompt)
        return {"evaluate": response.evaluate, "feedback": response.feedback, "feedback_history": response.feedback}

    # optimizing the content using the llm
    def optimize(state: studystate):
        prompt = f"optimize the content on topic {state['topic']} , material {state['material']}  and on the basis of the given feedback {state['feedback']} improve the material as and when required."

        response = optimized_model.invoke(prompt).content
        iteratation = state['iteration'] + 1

        return {"material": response, "iteration": iteratation, "material_history": [response]}

    def route_evaluation(state: studystate):
        if state['evaluate'] == 'approved' and state['iteration'] >= state["max_iteration"]:
            return "approved"
        else:
            return "not_approved"

    # ---------------------------------------------------------------------
   
    # uploading files
    st.subheader("Upload Your File Below👇")
    st.caption("✅Make sure The File Must Be Docx File")
    uploading_file = st.file_uploader("choose the file")
  
    final_text = []
    try:
        if uploading_file is not None:
            doc = Document(uploading_file)

            for para in doc.paragraphs:
                final_text.append(para.text)
        else:
            st.warning("please upload a file first 🗃️")
            
    except Exception as e :
        print(f"Error {e}")

    content = "\n".join(final_text)
    st.markdown(content)
    st.space("small")
    st.subheader("✍️Prompting")
    st.caption("you can edit the prompt as per your requirements")
    prompt = st.text_area(label="Enter the prompt:",value='''You are an experienced professional educator and tutor.
            Write a structured, theory-focused educational Study Material.
            Requirements:
            - Explain concepts clearly with descriptive theoretical depth, progressing from basics to key ideas,maintain the balance of theoretical and practical knowledge. 
            - Organize content into clear sections with headings.
            - the content generated must be static do not change the format.
            - In each section: (follow the points mentioned below) 
            • Provide paragraph-based theory explanations.
            • Include concise real-world examples with outcomes.(codes if required)
            • Add  5-10 descriptive practice questions  along with code outputs compulsory at the end of the content.
            - End with links for study resources relevant to the topic and beginner friendly websites(like w3schools,geeksforgeeks,etc.)
            - provide the links that exists.
            - Academic yet simple tone.
            - No visuals, emojis, filler, or repetition.
            - Be concise, accurate, and educational.
            - Ensure to make the use of seo keywords in the middle of the material.''',
            height=500,width='stretch')

    # creating and compiling the graph
    graph = StateGraph(studystate)

    # adding the node
    graph.add_node("material", study_material)
    graph.add_node("evaluating", evaluate_material)
    graph.add_node("optimize", optimize)

    # adding the edge
    graph.add_edge(START, "material")
    graph.add_conditional_edges("evaluating", route_evaluation, {
                                "approved": END, "not_approved": "optimize"})
    graph.add_edge("optimize", "evaluating")

    # compiling the graph
    workflow = graph.compile()
    
   
    initial_state = {
        "topic" :content,
        'prompt': prompt,
        "iteration": 1,
        "max_iteration": 3
    }

    generate = st.button(label="GENERATE FINAL MATERIAL",use_container_width=True)

    if generate:
        with st.spinner("⛷️Generating Response...."):
            time.sleep(5)
            final_output = workflow.invoke(initial_state)
            final_material = final_output['material']
            st.markdown(final_material)

            st.session_state['final_material'] = final_material

    # inserting the data into the file
    insert_in_file = st.button("INSERT MATERIAL INTO FILE", use_container_width=True)

    if insert_in_file:
        final_material_filename = uploading_file.name.replace(".docx", "")
        if "final_material" not in st.session_state:
            st.error("please generate final material first")

        else:
            try:
                doc = Document(f"{final_material_filename}_material.docx")

            except Exception as e:
                doc = Document()

            doc.add_paragraph("-------------FINAL MATERIAL---------------")
            doc.add_paragraph(st.session_state['final_material'])

            # save the doc
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            st.success("📝Data inserted successfully!!")

            # creating a download button to download the file
            st.download_button(
                label="DOWNLOAD FILE",
                data=bio.getvalue(),
                file_name=f"{final_material_filename}_material.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                icon=":material/download:",
                width='stretch'
            )

    html_content = st.button("GENERATE HTML", use_container_width=True)
    roadmap_button = st.button("GENERATE ROADMAP", use_container_width=True)
    outline_button = st.button("GENERATE OUTLINE", use_container_width=True)

    if html_content:
        st.switch_page("pages/convert_html.py")
    if roadmap_button:
        st.switch_page("pages/road_map.py")
    if outline_button:
        st.switch_page("pages/outline.py")


else:
    st.switch_page("password.py")

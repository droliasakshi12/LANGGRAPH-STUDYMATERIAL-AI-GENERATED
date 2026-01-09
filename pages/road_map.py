from dotenv import load_dotenv
import os
from typing import TypedDict, Literal, Annotated
from pydantic import BaseModel, Field
from langgraph.graph import StateGraph, START, END
from langchain_openai import ChatOpenAI
import operator
import streamlit as st
from docx import Document
from pages.model_selection import selecting_model
import io


if 'register' not in st.session_state:
    st.session_state['register'] = False


elif st.session_state['register'] == True:

    load_dotenv(dotenv_path=r"my_api_key.env")
    os.environ["OPENAI_API_KEY"] = os.getenv("openai_api") 
    os.environ["LANGCHAIN_TRACING_V2"] = os.getenv("LANGCHAIN_TRACING_V2") or st.secrets["LANGCHAIN_TRACING_V2"]
    os.environ["LANGCHAIN_ENDPOINT"] = os.getenv("LANGCHAIN_ENDPOINT") or st.secrets["LANGCHAIN_ENDPOINT"]
    os.environ["LANGCHAIN_API_KEY"] = os.getenv("langsmith_api") or st.secrets["langsmith_api"]
    os.environ["LANGCHAIN_PROJECT"] = os.getenv("LANGCHAIN_PROJECT") or st.secrets["LANGCHAIN_PROJECT"]

    # getting models
    # model = ChatOpenAI(model="gpt-4.1-nano-2025-04-14")
    evaluate_model = ChatOpenAI(model="gpt-3.5-turbo")
    optimized_model = ChatOpenAI(model="gpt-5-nano-2025-08-07")

    # -------------------MODEL SELECTION-----------------------------
    models = selecting_model()
    model = ChatOpenAI(model=models)

    # creating base model

    class evaluate(BaseModel):
        evaluate: Literal['approved', 'not_approved'] = Field(
            description="evaluating the content")
        feedback: str = Field(description="feedback for the content writen")

    struct_model = evaluate_model.with_structured_output(evaluate)

    # creating state model

    class roadstate(TypedDict):
        topic: str
        road_map: str = Field(description="generate a road map")
        evaluate: Literal['approved', 'not_approved']
        feedback: str
        iteration: int
        max_iteration: int
        content_history: str = Annotated[list[str], operator.add]
        feedback_history: str = Annotated[list[str], operator.add]

    # creating a function to generate the content

    def generate_roadmap(state: roadstate):
        prompt = f"""generate a short and comprehemsive road map on {state['topic']} as a study guide for students.
        replace the normal words with seo keywords where required.
        should not be in que/ans format make it in the format of roadmap ."""

        road_map = model.invoke(prompt).content

        return {"road_map": road_map, 'content_history': [road_map]}

    # evaluating the content generated

    def evaluating(state: roadstate):
        prompt = f"""you have to evaluate the content generated {state['road_map']} on the basis of this you have to check that the content generated is to be approved or not."""

        evaluating = struct_model.invoke(prompt).content

        return {"evaluate": evaluating, 'feedback': evaluating.feedback, 'feedback_history': evaluating.feedback}

    # optimizing the content generated

    def optimizing(state: roadstate):
        prompt = f"""improve the material generated and on the basics of the feedback check if the material generated is approved or not.
        here are the topic :{state['topic']} , content:{state["road_map"]} and the feedback :{state['feedback']}
        """

        optimize = optimized_model.invoke(prompt)
        iteration = state['iteration'] + 1

        return {'road_map': optimize, 'iteration': iteration, 'content_history': [optimize]}

    def evaluate_route(state: roadstate):
        if state['evalaute'] == "approved" and state['iteration'] >= state['max_iteration']:
            return "approved"
        else:
            return "not_approved"

    # --------------------------------------------------------------
    # taking user input
    user_input = st.text_input(label="ENTER THE TOPIC",
                               placeholder="explain python programming")

    generate = st.button("GENERATE THE ROADMAP", use_container_width=True)

    # creating and compiling the graph
    graph = StateGraph(roadstate)

    # adding nodes
    graph.add_node("generate_roadmap", generate_roadmap)
    graph.add_node('evaluate_content', evaluating)
    graph.add_node("optimize", optimizing)

    # adding edges
    graph.add_edge(START, "generate_roadmap")
    graph.add_conditional_edges("evaluate_content", evaluate_route, {
        "approved": END, "not_approved": "optimize"})
    graph.add_edge("optimize", "evaluate_content")

    # initialising data
    initial_state_road_map = {
        "topic": user_input,
        "iteration": 1,
        "max_iteration": 3
    }

    # compiling the graph into workflow

    try:
        if generate:
            roadmap_workflow = graph.compile()
            final_output_road_map = roadmap_workflow.invoke(initial_state_road_map)
            roadmap_content = final_output_road_map['road_map']
            st.markdown(roadmap_content)

            st.session_state['roadmap_content'] = roadmap_content

        # inserting data into file
        file_name = initial_state_road_map['topic']
        insert_file = st.button("INSERT DATA INTO FILE",
                                use_container_width=True)

        if insert_file:
            if "roadmap_content" not in st.session_state:
                st.error("please generate the cotent first!!")
            else:
                try:
                    doc = Document(f"{file_name}_roadmap.docx")

                except Exception as e:
                    doc = Document()

                doc.add_paragraph(st.session_state["roadmap_content"])

                # save the doc
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                st.success("Data inserted successfully!!")

                # creating a download button to download the file
                st.download_button(
                    label="DOWNLOAD FILE",
                    data=bio.getvalue(),
                    file_name=f"{file_name}_roadmap.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    icon=":material/download:",
                )

    except Exception as e:
        st.error("error")

    outline_button = st.button(
        "GENERATE OUTLINE", use_container_width=True)
    final_material_button = st.button(
        "GENERATE MATERIAL", use_container_width=True)

    if outline_button:
        st.switch_page("pages/outline.py")

    if final_material_button:
        st.switch_page("pages/final_material.py")

else:
    st.switch_page("password.py")





